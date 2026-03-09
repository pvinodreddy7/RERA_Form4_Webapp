from flask import Flask, render_template, request, send_file
import io, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para_run(paragraph, text, bold=False, size=10, color=None, underline=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.underline = underline
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_paragraph(doc, text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10,
                  space_before=0, space_after=4, color=None, underline=False, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if text:
        para_run(p, text, bold=bold, size=size, color=color, underline=underline)
    return p

def cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]

def inr(v):
    """Prefix with ₹ if numeric-looking."""
    if not v or str(v).upper() in ('NIL', '', '0', '—', '-'):
        return str(v) if v else ''
    s = str(v).strip()
    if s.startswith('₹') or s.startswith('Rs'):
        return s
    return f'₹ {s}'

def pct(v):
    s = str(v).strip().rstrip('%')
    try:
        float(s)
        return f'{s}%'
    except:
        return v

# ─── Document Text Extraction (for AI parsing of previous certificates) ───────

def extract_docx_text(file_bytes):
    """Extract all text from a .docx in document order (paragraphs + tables)."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(t.text or '' for t in element.iter(f'{{{ns}}}t')).strip()
            if text:
                lines.append(text)
        elif tag == 'tbl':
            for tr in element.iter(f'{{{ns}}}tr'):
                cells = []
                for tc in tr.findall(f'{{{ns}}}tc'):
                    cell_text = ''.join(t.text or '' for t in tc.iter(f'{{{ns}}}t')).strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    lines.append(' | '.join(cells))
    return '\n'.join(lines)

# ─── Header ───────────────────────────────────────────────────────────────────

def add_rera_header(doc):
    """RERA title block at top of each page."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    para_run(p1, '[KARNATAKA REAL ESTATE REGULATORY AUTHORITY', bold=True, size=11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    para_run(p2, '(Real Estate (Regulation & Development) Rules, 2017)]', bold=True, size=11)

def add_firm_header(doc):
    """PVKR firm header used on CA signature pages."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    para_run(p, 'PVKR & Co LLP', bold=True, size=13, color=(31, 73, 125))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    para_run(p2, 'Chartered Accountants', bold=True, size=11, color=(31, 73, 125))
    # divider
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(6)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot); pPr.append(pBdr)

def add_ca_signature(doc, d, udin_key='udin1'):
    """CA signature block — kept together on page via keep_with_next."""
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.keep_with_next = True

    def sig_p(text, bold=False, kwn=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = kwn
        para_run(p, text, bold=bold, size=10)
        return p

    sig_p('Yours Faithfully,')
    sig_p(f'For {d.get("firm_name","PVKR & Co LLP")} ')
    sig_p('Chartered Accountants,')
    sig_p(f'FRN: {d.get("frn","")}')
    # Three blank lines for physical signature space
    for _ in range(3):
        e = doc.add_paragraph()
        e.paragraph_format.space_after = Pt(3)
        e.paragraph_format.keep_with_next = True
    sig_p(f'CA {d.get("ca_name","")}', bold=True)
    sig_p(d.get('ca_designation', 'Partner'))
    sig_p(f'Membership Number: {d.get("membership_no","")}')
    e2 = doc.add_paragraph()
    e2.paragraph_format.space_after = Pt(3)
    e2.paragraph_format.keep_with_next = True
    sig_p(f'Date: {d.get("cert_date","")}')
    sig_p(f'Place: {d.get("place","Bangalore")}')
    udin = d.get(udin_key) or d.get('udin1', '')
    sig_p(f'UDIN: {udin}', kwn=False)   # last paragraph — no keep_with_next

# ─── Cost Table ───────────────────────────────────────────────────────────────

def build_cost_table(doc, d):
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'

    W = [Inches(0.4), Inches(3.9), Inches(1.3), Inches(1.3)]

    def hdr_row():
        row = tbl.add_row(); cells = row.cells
        for c in cells: set_bg(c, '1F497D')
        texts = ['Sl. No', 'Particulars', 'Estimated Amt in Rs.', 'Incurred Amt in Rs.']
        aligns = [WD_ALIGN_PARAGRAPH.CENTER]*4
        for i, (t, a) in enumerate(zip(texts, aligns)):
            cells[i].width = W[i]
            cell_para(cells[i], t, bold=True, size=9, align=a, color=(255,255,255))

    def data_row(sl, particulars, est, inc, bg=None, bold=False, indent=0):
        row = tbl.add_row(); cells = row.cells
        if bg:
            for c in cells: set_bg(c, bg)
        for i, w in enumerate(W): cells[i].width = w
        if sl:
            cell_para(cells[0], str(sl), bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if indent:
            p.paragraph_format.left_indent = Inches(indent * 0.15)
        run = p.add_run(particulars)
        run.bold = bold; run.font.size = Pt(9)
        cell_para(cells[2], inr(est) if est not in (None,'') else '', bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_para(cells[3], inr(inc) if inc not in (None,'') else '', bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        return row

    def section_row(sl, text):
        row = tbl.add_row(); cells = row.cells
        set_bg(cells[0], 'E8F0FB'); set_bg(cells[1], 'E8F0FB')
        set_bg(cells[2], 'E8F0FB'); set_bg(cells[3], 'E8F0FB')
        for i, w in enumerate(W): cells[i].width = w
        if sl: cell_para(cells[0], str(sl), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(cells[1], text, bold=True, size=9)

    def subtotal_row(label, est, inc):
        r = data_row('', label, est, inc, bg='D5E3F7', bold=True)

    def single_val_row(label, val, bg='EAF7EE', bold=True, colspan=False):
        row = tbl.add_row(); cells = row.cells
        for c in cells: set_bg(c, bg)
        for i, w in enumerate(W): cells[i].width = w
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        run = p.add_run(label); run.bold = bold; run.font.size = Pt(9)
        # Merge cols 2 and 3
        merged = cells[2].merge(cells[3])
        cell_para(merged, inr(val) if val not in (None,'') else '', bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    hdr_row()

    # Land Cost
    section_row('1', 'Land Cost')
    data_row('', 'a) Cost incurred for acquisition of ownership and title of the land parcels (outright purchase / lease etc.)',
             d.get('lc_a_est',''), d.get('lc_a_inc',''), indent=1)
    data_row('', 'b) Amount paid for Acquisition / purchase of TDR (if any)',
             d.get('lc_b_est',''), d.get('lc_b_inc',''), indent=1)
    data_row('', 'c) Amount paid to the Competent Authority for project approval, NOCs, stamp duty, transfer charges, registration charges, conversion charges, taxes, statutory payments to State and Central Government',
             d.get('lc_c_est',''), d.get('lc_c_inc',''), indent=1)
    subtotal_row('Sub – Total Land Cost', d.get('lc_total_est',''), d.get('lc_total_inc',''))

    # Development Cost
    section_row('', 'Development Cost / Cost of Construction')
    data_row('', '(i) Estimated Cost of Construction as certified by Engineer',
             d.get('dc_i_est',''), '', indent=1)
    data_row('', '(ii) Actual Cost of construction incurred as per books of accounts as verified by CA\n(Note: for adding to total cost of construction incurred, minimum of (i) or (ii) is to be considered)',
             '', d.get('dc_ii_inc',''), indent=1)
    data_row('', '(iii) On-site expenditure for development of entire project excluding cost of construction (salaries, consultants fees, site overheads, development works, cost of services including water, electricity, sewerage, drainage, layout roads etc., cost of machineries and equipment, consumables etc.)',
             d.get('dc_iii_est',''), d.get('dc_iii_inc',''), indent=1)
    data_row('', 'Payment of Taxes, cess, fees, charges, premiums, interest etc., to any statutory Authority',
             d.get('dc_tax_est',''), d.get('dc_tax_inc',''), indent=1)
    data_row('', 'Interest payable to financial institutions, scheduled banks, NBFCs or money lenders on construction funding',
             d.get('dc_int_est',''), d.get('dc_int_inc',''), indent=1)
    subtotal_row('Sub – Total Development Cost', d.get('dc_total_est',''), d.get('dc_total_inc',''))

    # Summary rows
    single_val_row('Total Estimated Cost of the Real Estate Project [Land Estimated + Development Estimated]',
                   d.get('total_est',''), bg='D5E3F7')
    single_val_row('Total Cost Incurred of the Real Estate Project [Land Incurred + Development Incurred]',
                   d.get('total_inc',''), bg='D5E3F7')
    single_val_row(f'Percentage of completion of construction work (as per Project Architect\'s Certificate)',
                   pct(d.get('arch_pct','')), bg='FAFCFF', bold=False)
    single_val_row('Proportion of the Cost incurred on Land Cost to the Total Estimated Cost',
                   d.get('prop_land',''), bg='EAF7EE')
    single_val_row('Proportion of the Cost incurred on Construction Cost to the Total Estimated Cost',
                   d.get('prop_dev',''), bg='EAF7EE')
    single_val_row('Total percentage of completion of construction as per CA (Proportion Land + Proportion Dev)',
                   d.get('ca_pct',''), bg='EAF7EE')
    single_val_row('Amount which can be withdrawn from Designated Account (Total Estimated Cost × Percentage of completion of construction) [lower of Architect % & CA %]',
                   d.get('withdrawable',''), bg='D5F0E8', bold=True)
    single_val_row('Less: Amount withdrawn till date of this certificate as per the Books of Accounts and Bank Statement',
                   d.get('withdrawn',''), bg='FAFCFF', bold=False)
    single_val_row('Net Amount which can be withdrawn from the Designated Bank Account under this certificate',
                   d.get('net_withdrawable',''), bg='D5F0E8', bold=True)

    # ── Helpers for label/value rows (used in sections 2 & 3) ──────────────────
    def lv_row(label, value, bg='FAFCFF', label_bold=False, indent=0.15):
        """Label in col1, value in merged cols 2+3."""
        row = tbl.add_row(); cells = row.cells
        for i, w in enumerate(W): cells[i].width = w
        for i in range(1, 4): set_bg(cells[i], bg)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.bold = label_bold
        merged = cells[2].merge(cells[3])
        set_bg(merged, bg)
        cell_para(merged, str(value) if value else '', size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    def b_sh_row(text, bg='EEF3FA'):
        """Sub-header spanning cols 1+2+3."""
        row = tbl.add_row(); cells = row.cells
        for i, w in enumerate(W): cells[i].width = w
        for i in range(4): set_bg(cells[i], bg)
        merged = cells[1].merge(cells[2]).merge(cells[3])
        p = merged.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.05)
        run = p.add_run(text); run.bold = True; run.font.size = Pt(9)

    # Section 2: Borrowings / Mortgages (supports multiple entries)
    section_row('2', 'Borrowings / Mortgage Details (If Applicable)')

    borrowings = d.get('borrowings', [])
    # backward compat: legacy single-entry fields
    if not borrowings and d.get('lender'):
        borrowings = [{'lender': d.get('lender',''), 'amt_disbursed': d.get('amt_disbursed',''),
                       'pending_disb': d.get('pending_disb',''), 'amt_repay': d.get('amt_repay','')}]

    b_sh_row('A. Borrowing Details')
    if borrowings:
        for idx, b in enumerate(borrowings, 1):
            pfx = f'({idx}) ' if len(borrowings) > 1 else ''
            lv_row(pfx + 'Name of the Lender', b.get('lender', '') or 'NIL')
            lv_row(pfx + 'Amount Disbursed', inr(b.get('amt_disbursed', '')))
            lv_row(pfx + 'Amount Pending for Disbursement from Lender', inr(b.get('pending_disb', '')))
            lv_row(pfx + 'Amount to be Repaid to Lender', inr(b.get('amt_repay', '')))
    else:
        lv_row('Not Applicable', 'NIL')

    mortgages = d.get('mortgages', [])
    # backward compat
    if not mortgages and d.get('mortgaged_to'):
        mortgages = [{'mortgaged_to': d.get('mortgaged_to',''), 'm_amt_disbursed': d.get('m_amt_disbursed',''),
                      'm_pending': d.get('m_pending',''), 'm_repay': d.get('m_repay','')}]

    b_sh_row('B. Mortgage Details')
    if mortgages:
        for idx, m in enumerate(mortgages, 1):
            pfx = f'({idx}) ' if len(mortgages) > 1 else ''
            lv_row(pfx + 'Mortgaged to (Name of Entity / Institution)', m.get('mortgaged_to', '') or 'NIL')
            lv_row(pfx + 'Amount Disbursed', inr(m.get('m_amt_disbursed', '')))
            lv_row(pfx + 'Amount Pending for Disbursement', inr(m.get('m_pending', '')))
            lv_row(pfx + 'Amount to be Repaid to Lender', inr(m.get('m_repay', '')))
    else:
        lv_row('Not Applicable', 'NIL')

    # Section 3: Bank Transactions
    section_row('3', 'Details of transactions in the designated RERA Bank Account (Including New Account)')
    lv_row('Total number of units booked', d.get('units_booked',''), bg='F7FAFF')
    lv_row('Total amount realized from sale of units during the quarter', inr(d.get('total_realized','')), bg='F7FAFF')
    lv_row('Total amount deposited into the bank out of sale proceeds during the quarter', inr(d.get('total_deposited','')), bg='F7FAFF')
    lv_row('% Of Deposit made', d.get('pct_deposited',''), bg='EAF7EE', label_bold=True)

    # Quarterly reconciliation
    b_sh_row('Reconciliation for the Quarter')
    lv_row('Opening Balance Date', d.get('q_open_date',''))
    lv_row('Opening Balance as per bank statement (INR)', inr(d.get('q_open_bal','')))
    lv_row('Deposits during the Quarter on account of sales (INR)', inr(d.get('q_dep_sales','')))
    lv_row('Other Deposits made (If any)', inr(d.get('q_dep_other','')))
    lv_row('Withdrawals during the Quarter from sale proceeds (INR)', inr(d.get('q_wdl_sales','')))
    lv_row('Other withdrawals made (if any)', inr(d.get('q_wdl_other','')))
    lv_row('Closing Balance as per bank statement (INR)', inr(d.get('q_close_bal','')), bg='D5E3F7', label_bold=True)
    lv_row('Closing Balance Date', d.get('q_close_date',''))

    # Cumulative reconciliation
    b_sh_row('Cumulative Reconciliation from beginning of project (till end of quarter)')
    lv_row('Opening balance of the account (INR)', inr(d.get('c_open_bal','')))
    lv_row('Total Deposits made from sale proceeds (including old RERA Account) (INR)', inr(d.get('c_dep_sales','')))
    lv_row('Total deposits made other than sale proceeds (if any) (INR)', inr(d.get('c_dep_other','')))
    lv_row('Total Withdrawals made from sale proceeds including Sales returns (INR)', inr(d.get('c_wdl_sales','')))
    lv_row('Total withdrawals made other than those from sale proceeds (if any) (INR)', inr(d.get('c_wdl_other','')))
    lv_row('Closing balance for the current quarter (INR)', inr(d.get('c_close_bal','')), bg='D5E3F7', label_bold=True)

    return tbl


# ─── Additional Info Table (Table 1 in original) ──────────────────────────────

def build_receivables_table(doc, d):
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'
    W = [Inches(0.4), Inches(4.6), Inches(1.8)]

    def hdr():
        row = tbl.add_row(); cells = row.cells
        for c in cells: set_bg(c, '1F497D')
        for i, w in enumerate(W): cells[i].width = w
        cell_para(cells[0], 'Sl.', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))
        cell_para(cells[1], 'Particulars', bold=True, size=9, color=(255,255,255))
        cell_para(cells[2], 'Amount (₹)', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))

    def row_add(sl, label, val, bg='FAFCFF', bold=False):
        row = tbl.add_row(); cells = row.cells
        for c in cells: set_bg(c, bg)
        for i, w in enumerate(W): cells[i].width = w
        if sl: cell_para(cells[0], str(sl), bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(cells[1], label, bold=bold, size=9)
        cell_para(cells[2], inr(val) if val not in (None,'') else '', bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    hdr()
    row_add('', 'Estimated Balance Cost to Complete the Real Estate Project (Total Estimated Cost − Total Cost Incurred)',
            d.get('bal_cost',''), bg='EAF7EE', bold=True)
    row_add('', 'Balance number of receivables from sold apartments as per Annexure A to this certificate (as certified by CA as verified from records and books of Accounts)',
            d.get('bal_receivables',''))
    row_add('3.', '(i) Balance Unsold area (To be certified by Management and verified by CA from records and books of accounts)',
            f"{d.get('unsold_area','')} Sq Mts")
    row_add('', '(ii) Estimated amount of sales proceeds in respect of unsold apartments (calculated as per ASR × unsold area as on date of certificate) as per Annexure A',
            d.get('unsold_asr_total',''))
    row_add('4.', 'Estimated receivables of ongoing project. Sum of 2 + 3(ii)',
            d.get('est_receivables',''), bg='D5E3F7', bold=True)
    row_add('5.', f'Amount to be deposited in Designated Account — 70% or 100%\n(If Sl.4 > Sl.1 → 70%; If Sl.4 ≤ Sl.1 → 100%)',
            d.get('amt_to_deposit',''), bg='D5F0E8', bold=True)
    return tbl


# ─── Main document generator ──────────────────────────────────────────────────

def generate_form4(d):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)

    # ── PAGE 1: Certificate ───────────────────────────────────────────────────
    add_rera_header(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    para_run(p, 'FORM-4', bold=True, size=12, underline=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    para_run(p2, "CHARTERED ACCOUNTANT'S CERTIFICATE", bold=True, size=11)

    # Date right-aligned
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pd.paragraph_format.space_after = Pt(4)
    para_run(pd, f'Date: {d.get("cert_date","")}', size=10)

    # Fields
    def field(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        para_run(p, f'{label}\t: {value}', size=10)

    field('KRERA Registration Number', d.get('krera_reg',''))
    field('Project Name    \t\t', d.get('project_name',''))
    field('Promoter Name\t\t', d.get('promoter_name',''))
    cost_str = f'₹ {d.get("project_cost","")}/-'
    if d.get('project_cost_words'):
        cost_str += f' ({d.get("project_cost_words","")})'
    field('Cost of Real Estate Project\t', cost_str)
    if d.get('quarter_label'):
        field('Quarter\t\t\t\t', d.get('quarter_label',''))
    field('Quarter End Date\t\t', d.get('quarter_end',''))

    doc.add_paragraph()

    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    para_run(body1, 'This Certificate is issued in accordance with the provisions of the Real Estate (Regulation and Development) Act, 2016 read with the Karnataka Real Estate (Regulation and Development Rules, 2017.', size=10)

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    para_run(body2, f'The Promoter in compliance with section 4(2)(l)(D), of the Real Estate (Regulation and Development) Act, 2016 has deposited 70% of the amounts received from the allottees of this project (refer observations) in the following account:', size=10)

    # Three bank accounts
    def bank_block(holder, krbad, acno, bank, ifsc, branch):
        for label, val in [
            ('Name of the Account Holder', holder),
            ('Name of the Designated bank account as per KRERA', krbad),
            ('Designated Account Number', acno),
            ('Bank Name\t\t\t', bank),
            ('IFSC Code\t\t\t', ifsc),
            ('Branch Name\t\t\t', branch),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.3)
            para_run(p, f'{label}\t: {val}', size=10)

    doc.add_paragraph()
    bank_block(d.get('b1_holder',''), d.get('b1_krbad',''), d.get('b1_acno',''),
               d.get('b1_bank',''), d.get('b1_ifsc',''), d.get('b1_branch',''))
    doc.add_paragraph()
    bank_block(d.get('b2_holder',''), d.get('b2_krbad',''), d.get('b2_acno',''),
               d.get('b2_bank',''), d.get('b2_ifsc',''), d.get('b2_branch',''))
    doc.add_paragraph()
    bank_block(d.get('b3_holder',''), d.get('b3_krbad',''), d.get('b3_acno',''),
               d.get('b3_bank',''), d.get('b3_ifsc',''), d.get('b3_branch',''))

    doc.add_paragraph()
    cert_detail = doc.add_paragraph()
    cert_detail.paragraph_format.space_after = Pt(6)
    para_run(cert_detail,
        f'This certificate is being issued for the project {d.get("project_name","")} '
        f'with RERA Registration No. : {d.get("krera_reg","")} for the quarter ended '
        f'{d.get("quarter_end","")} in compliance of the provisions of section 4(2)(l)(D) of the Act '
        f'and based on the records and documents produced before me and explanations provided to me '
        f'by the Management of the Company.', size=10)

    # Observations
    obs = d.get('observations','')
    obs_lines = [l.strip() for l in obs.split('\n') if l.strip()] if obs else []
    if obs_lines:
        doc.add_paragraph()
        pobs = doc.add_paragraph()
        pobs.paragraph_format.keep_with_next = True
        para_run(pobs, 'Qualification/Observations:', bold=True, size=10)
        doc.add_paragraph()
        for i, line in enumerate(obs_lines):
            pl = doc.add_paragraph()
            pl.paragraph_format.left_indent = Inches(0.3)
            pl.paragraph_format.space_after = Pt(3)
            # Last obs line keeps with the signature spacer
            pl.paragraph_format.keep_with_next = (i == len(obs_lines) - 1)
            para_run(pl, line, size=10)
    else:
        # No obs — cert_detail should pull signature with it
        cert_detail.paragraph_format.keep_with_next = True

    add_ca_signature(doc, d, 'udin1')

    # ── PAGE 2: Cost Table (Additional Information) ───────────────────────────
    doc.add_page_break()
    add_rera_header(doc)
    doc.add_paragraph()
    p_add = doc.add_paragraph()
    para_run(p_add, '(Additional Information for Projects)', bold=True, size=11,
             color=(31, 73, 125))
    p_add.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    build_cost_table(doc, d)

    # ── PAGE 3: Receivables / Additional Info table ───────────────────────────
    doc.add_page_break()
    add_rera_header(doc)
    doc.add_paragraph()

    build_receivables_table(doc, d)

    # ── PAGE 4: Annexure A — Sold Inventory ───────────────────────────────────
    sold = d.get('sold_inventory', [])
    if sold:
        doc.add_page_break()
        add_rera_header(doc)
        doc.add_paragraph()
        pan = doc.add_paragraph()
        para_run(pan, 'Annexure A', bold=True, size=12, underline=True)
        p_st = doc.add_paragraph()
        para_run(p_st, 'Statement for calculation of Receivables from the Sales of the Ongoing Real Estate Project',
                 bold=True, size=10)
        doc.add_paragraph()
        p_si = doc.add_paragraph()
        para_run(p_si, 'Sold Inventory', bold=True, size=10, underline=True)
        doc.add_paragraph()

        stbl = doc.add_table(rows=1, cols=8)
        stbl.style = 'Table Grid'
        SW = [Inches(0.35), Inches(0.7), Inches(0.9), Inches(1.1), Inches(1.1), Inches(1.1), Inches(1.0), Inches(0.75)]
        hdrs = ['Sr. No.', 'Flat No.', 'Carpet Area\n(sq.mts.)', 'Unit Consideration\nas per Agreement (₹)',
                'Received\nAmount (₹)', 'Balance\nReceivable (₹)', 'Date of\nAgreement', 'Registered?\n(Yes/No)']
        hrow = stbl.rows[0]
        for i, (h, w) in enumerate(zip(hdrs, SW)):
            hrow.cells[i].width = w
            set_bg(hrow.cells[i], '1F497D')
            cell_para(hrow.cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))

        total_bal = 0
        for idx, item in enumerate(sold, 1):
            row = stbl.add_row()
            for i, w in enumerate(SW): row.cells[i].width = w
            vals = [str(idx), item.get('flat_no',''), item.get('carpet_area',''),
                    item.get('agreement_price',''), item.get('received',''),
                    item.get('balance',''), item.get('date',''), item.get('registered','NO')]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
                      WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
            for i, (v, a) in enumerate(zip(vals, aligns)):
                if idx % 2 == 0: set_bg(row.cells[i], 'F5F8FD')
                cell_para(row.cells[i], v, size=8, align=a)
            try:
                bal_str = item.get('balance','0').replace(',','')
                total_bal += float(bal_str) if bal_str else 0
            except: pass

        # Total row
        trow = stbl.add_row()
        for i, w in enumerate(SW): trow.cells[i].width = w
        set_bg(trow.cells[5], 'D5E3F7')
        merged = trow.cells[0].merge(trow.cells[1]).merge(trow.cells[2]).merge(trow.cells[3]).merge(trow.cells[4])
        cell_para(merged, 'Total Balance Receivable', bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_para(trow.cells[5], f'{total_bal:,.0f}', bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── PAGE 5: Unsold Inventory ───────────────────────────────────────────────
    doc.add_page_break()
    add_rera_header(doc)
    doc.add_paragraph()

    p_ui = doc.add_paragraph()
    para_run(p_ui, 'Unsold Inventory Valuation', bold=True, size=11, underline=True)
    doc.add_paragraph()

    p_rr = doc.add_paragraph()
    para_run(p_rr,
        f'Ready Reckoner Rate as on the date of Certificate of the Residential/Commercial premises '
        f'Rs.{d.get("asr_rate","")} per sq. mts.', size=10)
    doc.add_paragraph()

    utbl = doc.add_table(rows=1, cols=5)
    utbl.style = 'Table Grid'
    UW = [Inches(0.5), Inches(1.1), Inches(1.5), Inches(2.0), Inches(1.7)]
    u_hdrs = ['Sr. No.', 'Total Unsold\nFlats', 'Carpet Area\n(in sq.mts.) (A)',
              'Unit Consideration as per Ready Reckoner Rate\n(ASR) (B) in INR', 'Total ASR\nconsideration (A*B)']
    uhrow = utbl.rows[0]
    for i, (h, w) in enumerate(zip(u_hdrs, UW)):
        uhrow.cells[i].width = w
        set_bg(uhrow.cells[i], '1F497D')
        cell_para(uhrow.cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))

    urow = utbl.add_row()
    for i, w in enumerate(UW): urow.cells[i].width = w
    u_vals = ['1', d.get('unsold_flats',''), d.get('unsold_total_area',''),
              d.get('asr_rate',''), d.get('unsold_total_asr','')]
    u_aligns = [WD_ALIGN_PARAGRAPH.CENTER]*3 + [WD_ALIGN_PARAGRAPH.RIGHT]*2
    for i, (v, a) in enumerate(zip(u_vals, u_aligns)):
        cell_para(urow.cells[i], v, size=9, align=a)

    doc.add_paragraph()

    # Certification text for Annexure — keep_with_next so signature follows on same page
    pcert = doc.add_paragraph()
    pcert.paragraph_format.space_after = Pt(4)
    pcert.paragraph_format.keep_with_next = True
    para_run(pcert,
        f'This certificate is being issued for RERA compliance for {d.get("promoter_name","")} '
        f'and is based on the records and documents produced before me and explanations provided to me '
        f'by the Management of the Company.', size=10)

    udin2_key = 'udin2' if d.get('udin2') else 'udin1'
    add_ca_signature(doc, d, udin2_key)

    return doc


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()
        doc = generate_form4(data)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        project = (data.get('project_name','RERA') or 'RERA').replace(' ','_')
        quarter = (data.get('quarter_end','') or '').replace(' ','_').replace('/','_')
        return send_file(buf, as_attachment=True,
                         download_name=f'Form4_{project}_{quarter}.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        import traceback
        return {'error': str(e), 'trace': traceback.format_exc()}, 500


@app.route('/parse_certificate', methods=['POST'])
def parse_certificate():
    """Parse a previous-quarter Form 4 .docx using Claude API and return structured JSON."""
    try:
        import anthropic as _ant
        import json as _json
    except ImportError:
        return {'error': 'anthropic package not installed. Run: pip install anthropic'}, 500
    try:
        f = request.files.get('file')
        if not f:
            return {'error': 'No file uploaded'}, 400
        text = extract_docx_text(f.read())
        if not text.strip():
            return {'error': 'Could not extract any text from this document.'}, 400

        client = _ant.Anthropic()
        prompt = (
            "You are parsing a Karnataka RERA Form 4 Chartered Accountant's Certificate document.\n"
            "Extract the following fields and return ONLY a valid JSON object — no markdown, no explanation.\n"
            'Use "" for any field not found. Numeric amounts should use Indian comma formatting (e.g. "10,00,00,000").\n\n'
            "Fields to extract:\n"
            "  krera_reg          — KRERA registration number\n"
            "  project_name       — project name\n"
            "  promoter_name      — promoter / company name\n"
            "  project_cost       — total project cost (digits + commas only, no ₹)\n"
            "  project_cost_words — cost in words\n"
            "  place              — place of signing\n"
            "  b1_holder, b1_krbad, b1_acno, b1_bank, b1_ifsc, b1_branch  — 100% Collection Account\n"
            "  b2_holder, b2_krbad, b2_acno, b2_bank, b2_ifsc, b2_branch  — 70% Designated Account\n"
            "  b3_holder, b3_krbad, b3_acno, b3_bank, b3_ifsc, b3_branch  — 30% Account\n"
            "  firm_name, frn, ca_name, ca_designation, membership_no\n"
            "  lc_a_est, lc_a_inc  — land acquisition cost (estimated, incurred)\n"
            "  lc_b_est, lc_b_inc  — TDR cost (estimated, incurred)\n"
            "  lc_c_est, lc_c_inc  — statutory costs (estimated, incurred)\n"
            "  dc_i_est            — construction cost estimated (Engineer certificate)\n"
            "  dc_ii_inc           — actual construction cost incurred\n"
            "  dc_iii_est, dc_iii_inc — on-site expenditure\n"
            "  dc_tax_est, dc_tax_inc — taxes/cess\n"
            "  dc_int_est, dc_int_inc — interest on construction finance\n"
            "  arch_pct            — architect completion percentage (number only)\n"
            "  withdrawn           — amount withdrawn till date\n"
            "  asr_rate            — ready reckoner rate per sq.mt.\n"
            "  unsold_flats        — number of unsold flats\n"
            "  unsold_total_area   — unsold carpet area in sq.mts.\n"
            "  total_saleable_area — total saleable area if mentioned\n\n"
            f"Document text:\n{text[:10000]}"
        )

        msg = client.messages.create(
            model='claude-3-haiku-20240307',
            max_tokens=2048,
            messages=[{'role': 'user', 'content': prompt}]
        )
        result = msg.content[0].text.strip()
        # Strip markdown fences if the model adds them
        if result.startswith('```'):
            result = '\n'.join(result.split('\n')[1:]).rsplit('```', 1)[0].strip()
        return _json.loads(result)

    except Exception as e:
        import traceback
        return {'error': str(e), 'trace': traceback.format_exc()}, 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5050))
    app.run(debug=True, port=port)
